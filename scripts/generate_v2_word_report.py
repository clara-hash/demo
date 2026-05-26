#!/usr/bin/env python3
"""Generate V2 Word report including V7 Neural IRL results."""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = "outputs"

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)
    return table

def add_fig(doc, path, caption, w=5.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(caption); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs: r.font.size = Pt(9); r.italic = True

def build():
    doc = Document()

    # Title
    t = doc.add_heading('基于nuPlan/Las Vegas地图的无人车IRL路径规划实验报告 (v2)', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    doc.add_heading('摘要', level=1)
    doc.add_paragraph(
        '本报告记录了在nuPlan Las Vegas真实城市地图数据上，基于逆强化学习（IRL）进行自动驾驶'
        '路径规划的完整实验迭代过程。实验从基础Goal-conditioned IRL（V3）开始，逐步演进到'
        '神经网络奖励模型（V7），在真实城市道路语义地图和专家轨迹数据上验证了从专家示范中'
        '学习驾驶偏好的能力。'
    )
    doc.add_paragraph('关键词：逆强化学习，路径规划，自动驾驶，nuPlan，最大熵模型，神经网络奖励').italic = True

    # ── Ch1 ──
    doc.add_heading('1. 实验背景与目标', level=1)
    doc.add_heading('1.1 问题定义', level=2)
    doc.add_paragraph(
        '给定城市道路语义地图和人类专家驾驶轨迹，通过逆强化学习从专家示范中学习奖励函数，'
        '使得在给定起点和目标点时，规划出的路径能够体现专家驾驶偏好（可行驶区域行驶、保持安全距离、高效抵达目标等）。'
    )
    doc.add_heading('1.2 数据基础', level=2)
    add_table(doc, ['项目', '内容'], [
        ['地图来源', 'nuPlan Las Vegas Strip HD Map'],
        ['地图格式', 'GPKG → 栅格化 (240×187 grids)'],
        ['语义层', '0=无效区, 1=可行驶区, 2=车道, 3=连接区, 4=交叉口'],
        ['有效栅格', '10,183 cells'],
        ['专家轨迹', '12条 nuPlan ego vehicle记录'],
    ])
    doc.add_paragraph()

    doc.add_heading('1.3 最大熵逆强化学习原理', level=2)
    doc.add_paragraph(
        '奖励函数定义为特征线性加权：R(s) = θᵀφ(s)。最大熵模型下轨迹分布为：'
        'P(τ|θ) = exp(R(τ)) / Z(θ)。通过最大化专家轨迹对数似然学习参数θ。'
    )

    # ── Ch2: V3 ──
    doc.add_heading('2. V3：目标条件Goal-conditioned IRL（基线）', level=1)
    doc.add_paragraph('在语义地图上给定起点和目标点，通过候选路径特征评分选择最优路径，使用expert_density作为辅助特征。')
    add_table(doc, ['参数', '值'], [
        ['特征维度', '9'], ['训练样本', '43'], ['候选路径/样本', '18'],
        ['训练轮数', '220'], ['学习率', '0.06'], ['L2正则', '1e-4'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        '特征：drivable_ratio, lane_ratio, lane_connector_ratio, intersection_ratio, '
        'expert_density, goal_reached, heading_to_goal, smoothness, length_efficiency'
    )
    add_table(doc, ['指标', '值'], [
        ['初始NLL', '2.052'], ['最终NLL', '1.688'], ['IRL ADE', '2.600'],
        ['Shortest ADE', '1.315'], ['IRL成功率', '100%'],
        ['IRL路径长度', '36.014'], ['专家路径长度', '36.410'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        '学习权重：length_efficiency(1.338) > drivable_ratio(0.966) > intersection_ratio(0.953) > '
        'heading_to_goal(0.861) > expert_density(0.329) > smoothness(0.274)。'
        '依赖expert_density，缺少clearance和障碍物感知。'
    )

    # ── Ch3: V4 ──
    doc.add_heading('3. V4：Clearance-aware Goal-conditioned IRL', level=1)
    doc.add_paragraph(
        '引入static clearance特征，用EDT距离变换计算每个栅格到最近非行驶区域的距离，解决V3路径贴墙问题。'
        '移除lane_connector_ratio，新增clearance。'
    )
    add_table(doc, ['指标', 'V3', 'V4', '变化'], [
        ['最终NLL', '1.688', '1.662', '↓1.5%'],
        ['IRL ADE', '2.600', '1.919', '↓26.2%'],
        ['Shortest ADE', '1.315', '1.315', '-'],
        ['IRL路径长度', '36.014', '36.427', '↑1.1%'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        'Clearance是最有效的单点改进——ADE下降26.2%。路径长度微增1.1%，说明模型在安全-效率权衡中偏向安全。'
    )

    # ── Ch4: V5 ──
    doc.add_heading('4. V5：轻量级自主障碍感知IRL', level=1)
    doc.add_paragraph(
        '首次实现Train/Test按demo隔离；引入人工障碍物增强；移除expert_density；保存模型权重。'
        '但180样本太少，训练收敛不足，权重值域0-0.2，证明需要更大训练规模。'
    )
    add_table(doc, ['参数', 'V3/V4', 'V5'], [
        ['训练样本', '43', '180'], ['训练轮数', '220', '260'],
        ['学习率', '0.06', '0.035'], ['Train/Test分离', '按样本', '按demo隔离'],
        ['expert_density', '保留', '移除'],
    ])
    doc.add_paragraph()

    # ── Ch5: V6 Heavy ──
    doc.add_heading('5. V6 Heavy：大规模自主障碍感知线性IRL', level=1)
    doc.add_paragraph('大幅度提升规模：180 → 2,400 cases（+1233%）。900 epochs，24 candidates/case，7 obstacles/case。')
    add_table(doc, ['参数', 'V5', 'V6 Heavy'], [
        ['训练cases', '180', '2,400'], ['候选路径/case', '18', '24'],
        ['障碍物/case', '若干', '7'], ['训练轮数', '260', '900'],
        ['Batch size', '全量', '48'], ['有效候选集', '-', '2,371'],
        ['特征维度', '11', '12'], ['预处理耗时', '-', '30分钟'],
    ])
    doc.add_paragraph()

    doc.add_heading('5.1 Loss收敛', level=2)
    add_table(doc, ['轮数', '0', '100', '200', '300', '400', '500', '600', '700', '800', '899'],
              [['NLL', '2.664', '2.209', '2.057', '1.979', '1.932', '1.901', '1.879', '1.862', '1.848', '1.838']])
    doc.add_paragraph('Loss下降31.0%。')
    doc.add_paragraph()

    doc.add_heading('5.2 最终奖励权重', level=2)
    add_table(doc, ['特征', '权重', '方向', '解读'], [
        ['length_efficiency', '+19.738', '正向', '偏好高效路径（最强信号）'],
        ['static_clearance_min', '+7.715', '正向', '偏好安全（min>mean）'],
        ['heading_to_goal', '+6.488', '正向', '偏好朝向目标'],
        ['static_clearance_mean', '-15.369', '负向', '与min互补约束'],
        ['obstacle_clearance_mean', '-5.481', '负向', '统计偏差'],
        ['lane_ratio', '-4.069', '负向', '地图语义偏差'],
        ['smoothness', '-3.879', '负向', '候选生成偏见'],
    ])
    doc.add_paragraph()

    doc.add_heading('5.3 测试结果（30 cases）', level=2)
    add_table(doc, ['指标', 'IRL', 'Shortest'], [
        ['成功率', '100%', '100%'], ['平均路径长度', '133.85', '134.08'],
        ['障碍物违规', '0.0', '0.0'], ['地图违规', '0.0', '0.0'],
        ['平均转弯角', '0.0593', '0.0599'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        '规模效果显著，min clearance比mean更有信息量，完全摆脱expert_density。'
        '但线性奖励在简单场景与shortest baseline差异不大，需要非线性模型。'
    )

    # ── Ch6: V7 Neural ──
    doc.add_heading('6. V7 Neural Reward IRL', level=1)

    doc.add_heading('6.1 设计动机', level=2)
    doc.add_paragraph(
        'V6线性奖励R(s)=θᵀφ(s)表达能力有限。V7用MLP神经网络替代线性模型，'
        '从"特征权重学习"升级为"奖励函数学习"。采用选择式（Choice-based）学习方法：'
        '生成候选路径→teacher utility打分→MLP学习→Cross-entropy loss。'
    )

    doc.add_heading('6.2 模型架构', level=2)
    doc.add_paragraph(
        'Input(18) → Linear(256)+ReLU+Dropout(0.10) → Linear(256)+ReLU+Dropout(0.10) '
        '→ Linear(128)+ReLU → Linear(1) → scalar reward score'
    )

    doc.add_heading('6.5 Teacher Utility（修复后）', level=2)
    doc.add_paragraph(
        '第一版训练后发现障碍物间隙不足，第二版大幅提升了障碍物相关权重：'
    )
    add_table(doc, ['权重项', '第一版', '第二版（修复后）'], [
        ['obstacle_clearance_min', '3.4', '12.0'],
        ['obstacle_clearance_mean', '1.8', '5.0'],
        ['combined_clearance_min', '2.5', '8.0'],
        ['near_obstacle_penalty', '-6.5', '-15.0'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        '候选路径生成的障碍物硬约束从 min_obstacle_clearance=0~2 提升到 2~4 cells，'
        'A* 障碍物成本衰减范围从 /2.0 扩大到 /3.0，obstacle_weight 从 3~7 提升到 8~15。'
    )

    doc.add_heading('6.6 训练配置', level=2)
    add_table(doc, ['参数', '值'], [
        ['数据集', '3,000 cases × 10 candidates'],
        ['障碍物', '6-16椭圆随机障碍物/case'],
        ['起点-目标距离', '45-180 cells'],
        ['Train/Val/Test', '2100 / 450 / 450 (70/15/15)'],
        ['训练轮数', '80'], ['Batch size', '256'],
        ['学习率', '2e-4'], ['优化器', 'AdamW + CosineAnnealingLR'],
        ['设备', 'CPU (Jetson ARM)'], ['数据生成时间', '~75分钟'],
    ])
    doc.add_paragraph()

    doc.add_heading('6.7 训练结果', level=2)
    add_table(doc, ['阶段', 'Loss', 'Accuracy'], [
        ['Epoch 0 Train', '1.763', '31.3%'],
        ['Epoch 10 Val', '1.190', '61.3%'],
        ['Epoch 30 Val', '0.865', '73.6%'],
        ['Epoch 60 Val', '0.810', '77.3%'],
        ['Epoch 79 Val', '0.807', '77.8%'],
        ['Test (hold-out)', '0.868', '75.8%'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        'Loss下降54%。Choice Accuracy从随机水平10%提升到77.8%（val）、75.8%（test），'
        '较800 cases版本提升+11.6pp（test）。模型在约3/4情况下成功选中最优路径。'
    )

    doc.add_heading('6.8 自主规划测试（30个新障碍场景）', level=2)
    add_table(doc, ['指标', 'V7 Neural IRL', 'Shortest'], [
        ['成功率', '100%', '100%'],
        ['平均路径长度', '142.8', '134.5'],
        ['长度比均值', '1.070', '1.0'],
        ['平均障碍物间隙', '5.64', '-'],
        ['平均静态间隙', '1.22', '-'],
        ['obs_min ≥ 10.0 案例', '4个 (14.8/12.1/10.6/12.4)', '-'],
    ])
    doc.add_paragraph()

    doc.add_heading('6.9 训练规模影响：800 vs 3000 cases', level=2)
    add_table(doc, ['指标', '800 cases', '3000 cases', '变化'], [
        ['Val Accuracy', '62.5%', '77.8%', '↑+15.3pp'],
        ['Test Accuracy', '64.2%', '75.8%', '↑+11.6pp'],
        ['Val Loss', '1.104', '0.807', '↓27%'],
        ['平均障碍物间隙', '3.73', '5.64', '↑+51%'],
        ['obs_min ≤ 1.0', '17%', '6.7%', '↓61%'],
        ['obs_min ≥ 10.0', '0%', '13.3%', '新增'],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        '3000 cases训练带来质的飞跃：准确率突破75%，障碍物间隙突破5 cells，'
        '首次出现≥10 cells的极高安全间隙案例。'
    )

    doc.add_heading('6.10 V6 vs V7 对比', level=2)
    add_table(doc, ['维度', 'V6 Heavy (线性)', 'V7 (神经，3000 cases)'], [
        ['奖励函数', 'θᵀφ(s) 线性', 'MLP(φ) 非线性'],
        ['训练目标', 'MaxEnt特征匹配', 'Cross-Entropy选择'],
        ['测试行为', '路径≈shortest', '路径≠shortest（主动规避障碍物）'],
        ['障碍物间隙', '~1.44 (mix)', '~5.64 (obstacle min)'],
        ['路径变长代价', '~0%', '~6.9%'],
        ['训练规模', '2,400 cases', '3,000 cases'],
        ['选择准确率', '-', '75.8% (test)'],
    ])
    doc.add_paragraph()

    # ── Ch7: Summary ──
    doc.add_heading('7. 版本对比总结', level=1)
    add_table(doc, ['版本', '核心改进', '训练量', '特征', 'Loss变化', '关键指标'], [
        ['V3', 'Goal-conditioned基线', '43', '9', '2.052→1.688', 'ADE=2.600'],
        ['V4', '+Static Clearance', '43', '9', '2.052→1.662', 'ADE=1.919 (↓26%)'],
        ['V5', '+Obstacle, Train/Test分离', '180', '11', '-', '训练不足'],
        ['V6 Heavy', '规模扩大', '2,400', '12', '2.664→1.838', 'Success=100%'],
        ['V7', '神经网络奖励+障碍物规避', '3,000', '18→MLP', '1.763→0.807', 'Acc=75.8%, obs_min=5.64'],
    ])
    doc.add_paragraph()

    doc.add_heading('7.1 核心发现', level=2)
    findings = [
        'Clearance是最有效的单点改进：V3→V4 ADE下降26.2%',
        '规模决定质量：V5(180)权重接近零，V6(2400)学到有意义的权重',
        'min clearance比mean clearance更有信息量',
        '线性→神经是质的飞跃：V7首次让IRL路径系统性区分于shortest baseline',
        'Choice-based learning可行：65.8% accuracy在800 cases下已验证',
    ]
    for i, f in enumerate(findings):
        doc.add_paragraph(f'{i+1}. {f}')

    # ── Ch8: Conclusion ──
    doc.add_heading('8. 结论与展望', level=1)
    doc.add_heading('8.1 主要结论', level=2)
    for i, c in enumerate([
        '最大熵IRL能从nuPlan专家轨迹中学习有效驾驶偏好，系统性地验证了V3→V7的演进',
        'Clearance特征（静态间隙、障碍物间隙）是安全感知的关键桥梁',
        '神经网络奖励（V7）突破了线性表达的瓶颈，首次让IRL路径系统性地区分于shortest baseline',
        'Choice-based learning方法可行，即使在小规模下也能达到65.8%选择准确率',
    ]):
        doc.add_paragraph(f'{i+1}. {c}')

    doc.add_heading('8.2 后续工作', level=2)
    for f in [
        '扩大V7训练规模：800→10,000+ cases',
        '用nuPlan真实ego轨迹替代teacher utility作为标签',
        '探索Transformer/Attention架构建模路径级依赖',
        '闭环仿真评价（nuPlan simulator）',
        '引入交通规则约束（红绿灯、停止线、人行横道）',
    ]:
        doc.add_paragraph(f, style='List Bullet')

    # ── Appendix ──
    doc.add_heading('附录A：实验环境', level=1)
    add_table(doc, ['项目', '配置'], [
        ['硬件', 'NVIDIA Jetson (ARM aarch64, 61GiB RAM)'],
        ['OS', 'Ubuntu 20.04.6 LTS, L4T R35.4.1'],
        ['Python', '3.9.23 (Miniforge conda: nuplan_irl)'],
        ['依赖', 'numpy, scipy, matplotlib, pandas, shapely, PyTorch 2.8.0'],
        ['数据', 'nuPlan v1.1 Mini Split + Maps v1.0'],
        ['地图', 'Las Vegas Strip'],
    ])
    doc.add_paragraph()

    doc.add_heading('附录B：实验图表', level=1)

    # Insert figures
    add_fig(doc, 'outputs/report_all_versions_loss.png', '图1: V3/V4/V6/V7 训练Loss曲线对比（四象限）')
    doc.add_paragraph()
    add_fig(doc, 'outputs/report_v7_training.png', '图2: V7 训练Loss和Choice Accuracy曲线')
    doc.add_paragraph()
    add_fig(doc, 'outputs/report_v7_planning_summary.png', '图3: V7 24个测试case详细对比（路径长度、比率、间隙分布）')
    doc.add_paragraph()
    if os.path.exists('outputs/report_v7_scale_comparison.png'):
        add_fig(doc, 'outputs/report_v7_scale_comparison.png', '图4: V7 训练规模影响：800 vs 3000 cases')
        doc.add_paragraph()
    if os.path.exists('outputs/nuplan_irl_neural_reward_v7/planning_summary_v7.png'):
        add_fig(doc, 'outputs/nuplan_irl_neural_reward_v7/planning_summary_v7.png',
                '图6: V7 规划测试总体对比柱状图')
        doc.add_paragraph()

    # V7 test case examples
    fig_dir = 'outputs/nuplan_irl_neural_reward_v7/test_figures'
    if os.path.isdir(fig_dir):
        figs = sorted(os.listdir(fig_dir))[:4]
        for fn in figs:
            add_fig(doc, os.path.join(fig_dir, fn), f'V7测试case: {fn}', w=4.5)
            doc.add_paragraph()

    # Save
    out = os.path.join(OUT_DIR, 'IRL_实验报告_V3-V7_v2.docx')
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    build()
