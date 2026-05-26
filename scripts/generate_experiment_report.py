#!/usr/bin/env python3
"""Generate experiment report for IRL path planning V3-V6 on nuPlan Las Vegas map."""

import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

OUT_DIR = "outputs"
os.makedirs(OUT_DIR, exist_ok=True)

# ── Style helpers ──────────────────────────────────────────────
def set_cell_shading(cell, color):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_figure(doc, path, caption, width=5.5):
    """Insert a figure with caption."""
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.italic = True

# ── Generate comparison figure ──────────────────────────────────
def make_figures():
    """Generate comparison figures for the report."""

    # Figure 1: V3 vs V4 loss comparison
    fig, axes = plt.subplots(1, 3, figsize=(15, 4.5), dpi=140)

    # V3 loss (epoch 0, 10, 20, ..., 219 → 23 data points for 220 epochs)
    v3_loss = [
        2.052223, 2.016653, 1.984781, 1.956067, 1.930067, 1.906414,
        1.884800, 1.864967, 1.846699, 1.829813, 1.814151, 1.799582,
        1.785988, 1.773273, 1.761348, 1.750139, 1.739580, 1.729614,
        1.720188, 1.711258, 1.702783, 1.694728
    ]
    v3_epochs = np.arange(0, 220, 10)[:len(v3_loss)]
    axes[0].plot(v3_epochs, v3_loss, 'b-o', markersize=3, label='V3 NLL')
    axes[0].set_title('V3: Goal-conditioned IRL', fontsize=11, fontweight='bold')
    axes[0].set_xlabel('Epoch')
    axes[0].set_ylabel('Negative Log Likelihood')
    axes[0].grid(True, alpha=0.3)
    axes[0].text(0.95, 0.95, f'Final: {v3_loss[-1]:.3f}', transform=axes[0].transAxes,
                 ha='right', va='top', fontsize=9, color='blue')

    # V4 loss (epoch 0-219, step 10 → 23 points)
    v4_loss = [
        2.052223, 2.011472, 1.975561, 1.943679, 1.915179, 1.889537,
        1.866331, 1.845213, 1.825901, 1.808159, 1.791790, 1.776631,
        1.762541, 1.749402, 1.737113, 1.725587, 1.714748, 1.704532,
        1.694880, 1.685742, 1.677075, 1.668839, 1.661766
    ]
    v4_epochs = np.arange(0, 230, 10)[:len(v4_loss)]
    axes[1].plot(v4_epochs, v4_loss, 'g-s', markersize=3, label='V4 NLL')
    axes[1].set_title('V4: + Clearance Feature', fontsize=11, fontweight='bold')
    axes[1].set_xlabel('Epoch')
    axes[1].grid(True, alpha=0.3)
    axes[1].text(0.95, 0.95, f'Final: {v4_loss[-1]:.3f}', transform=axes[1].transAxes,
                 ha='right', va='top', fontsize=9, color='green')

    # V6 Heavy loss (31 points: every 30 epochs from 0 to 899)
    v6_loss = [
        2.664446, 2.564509, 2.390964, 2.285142, 2.208913, 2.150786,
        2.104984, 2.068127, 2.037867, 2.012751, 1.991498, 1.973348,
        1.957696, 1.944042, 1.932061, 1.921450, 1.911982, 1.903508,
        1.895839, 1.888876, 1.882512, 1.876680, 1.871313, 1.866351,
        1.861745, 1.857454, 1.852168, 1.847317, 1.843917, 1.840718,
        1.837792
    ]
    v6_epochs = np.arange(0, 930, 30)[:len(v6_loss)]
    axes[2].plot(v6_epochs, v6_loss, 'r-^', markersize=3, label='V6 Heavy NLL')
    axes[2].set_title('V6 Heavy: 2400 Cases, 900 Epochs', fontsize=11, fontweight='bold')
    axes[2].set_xlabel('Epoch')
    axes[2].grid(True, alpha=0.3)
    axes[2].text(0.95, 0.95, f'Final: {v6_loss[-1]:.3f}', transform=axes[2].transAxes,
                 ha='right', va='top', fontsize=9, color='red')

    fig.suptitle('Training Loss Convergence Across Versions', fontsize=13, fontweight='bold')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT_DIR, "report_loss_comparison.png"), dpi=160)
    plt.close(fig)

    # Figure 2: ADE comparison V3 vs V4
    fig, ax = plt.subplots(figsize=(6, 4), dpi=140)
    versions = ['V3', 'V4']
    ade_irl = [2.600, 1.919]
    ade_shortest = [1.315, 1.315]
    x = np.arange(len(versions))
    w = 0.35
    bars1 = ax.bar(x - w/2, ade_irl, w, label='IRL ADE', color='steelblue')
    bars2 = ax.bar(x + w/2, ade_shortest, w, label='Shortest ADE', color='lightcoral')
    for bar, val in zip(bars1, ade_irl):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f'{val:.3f}',
                ha='center', va='bottom', fontsize=10, fontweight='bold')
    for bar, val in zip(bars2, ade_shortest):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f'{val:.3f}',
                ha='center', va='bottom', fontsize=10)
    ax.set_ylabel('ADE (cells)')
    ax.set_title('V3 vs V4: Average Displacement Error', fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(versions)
    ax.legend()
    ax.grid(True, alpha=0.2, axis='y')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT_DIR, "report_ade_comparison.png"), dpi=160)
    plt.close(fig)

    # Figure 3: V6 Heavy reward weights
    fig, ax = plt.subplots(figsize=(8, 5), dpi=140)
    features = [
        'length_efficiency', 'static_clearance_min', 'heading_to_goal',
        'obstacle_clearance_min', 'lane_ratio', 'non_intersection_ratio',
        'obstacle_clearance_mean', 'smoothness', 'static_clearance_mean',
    ]
    weights = [19.738, 7.715, 6.488, -0.442, -4.069, -4.069, -5.481, -3.879, -15.369]
    colors = ['#2ca02c' if w > 0 else '#d62728' for w in weights]
    bars = ax.barh(features[::-1], weights[::-1], color=colors[::-1])
    ax.axvline(0, color='black', linewidth=0.5)
    ax.set_xlabel('Reward Weight')
    ax.set_title('V6 Heavy: Learned Reward Weights', fontweight='bold')
    for bar, val in zip(bars, weights[::-1]):
        ax.text(bar.get_width() + (0.3 if val >= 0 else -0.3),
                bar.get_y() + bar.get_height()/2,
                f'{val:.1f}', va='center',
                ha='left' if val >= 0 else 'right', fontsize=9)
    ax.grid(True, alpha=0.2, axis='x')
    fig.tight_layout()
    fig.savefig(os.path.join(OUT_DIR, "report_v6_weights.png"), dpi=160)
    plt.close(fig)

    print("Figures saved to outputs/")

# ── Build Word document ────────────────────────────────────────
def build_docx():
    doc = Document()

    # Title
    title = doc.add_heading('基于nuPlan/Las Vegas地图的无人车IRL路径规划实验报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    doc.add_heading('摘要', level=1)
    doc.add_paragraph(
        '本报告记录了在nuPlan Las Vegas真实城市地图数据上，基于最大熵逆强化学习'
        '（MaxEnt IRL）进行自动驾驶路径规划的实验迭代过程。实验共经历四个核心版本'
        '（V3-V6），从基础的Goal-conditioned IRL逐步演进到大规模障碍感知线性IRL，'
        '并设计了神经网络奖励模型的V7版本。在真实城市道路语义地图和专家轨迹数据上'
        '验证了IRL方法从专家示范中学习驾驶偏好的能力。'
    )
    doc.add_paragraph(
        '关键词：逆强化学习，路径规划，自动驾驶，nuPlan，最大熵模型'
    ).italic = True

    # ── Chapter 1: Background ──
    doc.add_heading('1. 实验背景与目标', level=1)

    doc.add_heading('1.1 问题定义', level=2)
    doc.add_paragraph(
        '给定城市道路语义地图和人类专家驾驶轨迹，通过逆强化学习从专家示范中学习奖励函数，'
        '使得在给定起点和目标点时，规划出的路径能够体现专家驾驶偏好：在可行驶区域行驶、'
        '保持安全距离、高效抵达目标等。'
    )

    doc.add_heading('1.2 数据基础', level=2)
    add_styled_table(doc,
        ['项目', '内容'],
        [
            ['地图来源', 'nuPlan Las Vegas Strip HD Map'],
            ['地图格式', 'GPKG → 栅格化 (240×187 grids)'],
            ['语义层', '0=无效区, 1=可行驶区, 2=车道, 3=车道连接区, 4=交叉口'],
            ['有效栅格', '10,183 cells'],
            ['车道栅格', '3,846 cells'],
            ['交叉口栅格', '6,313 cells'],
            ['专家轨迹', '12条nuPlan ego vehicle记录'],
            ['轨迹点数', '13–85 points/trajectory'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('1.3 核心算法：最大熵逆强化学习', level=2)
    doc.add_paragraph(
        '最大熵逆强化学习（Maximum Entropy IRL）假设专家轨迹服从最大熵分布。'
        '奖励函数定义为特征的线性加权：R(s) = θᵀφ(s)。在最大熵模型中，轨迹分布为：'
        'P(τ|θ) = exp(R(τ)) / Z(θ)。通过最大化专家轨迹的对数似然来学习参数θ：'
        'L(θ) = (1/|D|) Σ log P(τ|θ)。梯度计算通过匹配专家特征期望和模型特征期望来实现。'
    )

    # ── Chapter 2: V3 ──
    doc.add_heading('2. V3：目标条件Goal-conditioned IRL（基线版本）', level=1)

    doc.add_heading('2.1 设计思路', level=2)
    doc.add_paragraph(
        '在地图上给定起点和目标点，通过候选路径特征评分选择最优路径，'
        '使用专家轨迹密度作为辅助特征来指导路径选择。'
    )

    doc.add_heading('2.2 配置参数', level=2)
    add_styled_table(doc,
        ['参数', '值'],
        [
            ['特征维度', '9'],
            ['训练样本', '43'],
            ['候选路径/样本', '18'],
            ['训练轮数', '220'],
            ['学习率', '0.06'],
            ['L2正则', '1e-4'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('2.3 特征列表', level=2)
    add_styled_table(doc,
        ['特征', '含义'],
        [
            ['drivable_ratio', '可行驶区域占比'],
            ['lane_ratio', '车道区域占比'],
            ['lane_connector_ratio', '连接区占比'],
            ['intersection_ratio', '交叉口占比'],
            ['expert_density', '专家轨迹密度'],
            ['goal_reached', '到达目标点'],
            ['heading_to_goal', '朝向目标'],
            ['smoothness', '路径平滑度'],
            ['length_efficiency', '路径长度效率'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('2.4 训练结果', level=2)
    add_styled_table(doc,
        ['指标', '值'],
        [
            ['初始NLL', '2.052'],
            ['最终NLL', '1.688'],
            ['IRL ADE', '2.600'],
            ['Shortest ADE', '1.315'],
            ['IRL成功率', '100%'],
            ['IRL地图违规', '0'],
            ['IRL路径长度', '36.014'],
            ['专家路径长度', '36.410'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('2.5 学习的奖励权重', level=2)
    doc.add_paragraph(
        'length_efficiency(1.338) > drivable_ratio(0.966) > intersection_ratio(0.953) > '
        'heading_to_goal(0.861) > expert_density(0.329) > smoothness(0.274) > lane_ratio(0.013)'
    )

    doc.add_heading('2.6 分析', level=2)
    doc.add_paragraph(
        '奖励权重显示模型主要偏好高效（length_efficiency最高）且在可行驶区域的路径。'
        'expert_density有正向权重（0.329），说明模型依赖专家轨迹密度特征，泛化性受限。'
        '缺少static clearance特征，无法主动避开墙壁/边界。没有障碍物感知能力。'
    )

    # ── Chapter 3: V4 ──
    doc.add_heading('3. V4：Clearance-aware Goal-conditioned IRL', level=1)

    doc.add_heading('3.1 改进点', level=2)
    doc.add_paragraph(
        '在V3基础上引入static clearance（静态间隙）特征。用欧几里得距离变换（EDT）计算每个栅格'
        '到最近非行驶区域的距离，加入clearance作为路径点的平均静态间隙特征，以解决V3中路径贴墙的问题。'
        '特征变化：移除lane_connector_ratio，新增clearance（CLEARANCE_CAP=10cells, MIN=5cells）。'
    )

    doc.add_heading('3.2 训练结果对比', level=2)
    add_styled_table(doc,
        ['指标', 'V3', 'V4', '变化'],
        [
            ['最终NLL', '1.688', '1.662', '↓1.5%'],
            ['IRL ADE', '2.600', '1.919', '↓26.2%'],
            ['Shortest ADE', '1.315', '1.315', '-'],
            ['IRL成功率', '100%', '100%', '-'],
            ['IRL路径长度', '36.014', '36.427', '↑1.1%'],
            ['专家路径长度', '36.410', '36.410', '-'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('3.3 分析', level=2)
    doc.add_paragraph(
        'IRL ADE从2.600降到1.919，提升26.2%，说明clearance特征显著改善了路径质量。'
        'clearance权重0.779排在中间位置，说明它是有区分力的特征。'
        '路径长度从36.014增到36.427（增加1.1%），说明模型在做安全-效率权衡时略微偏向了安全。'
        '但该版本仍依赖expert_density特征，泛化性有限。'
    )

    # ── Chapter 4: V5 ──
    doc.add_heading('4. V5：轻量级自主障碍感知IRL', level=1)

    doc.add_heading('4.1 架构变化', level=2)
    doc.add_paragraph(
        '本次为架构重大变化：首次实现训练/测试分离（最后一条demo作为hold-out测试集）；'
        '引入人工障碍物增强（训练时在地图上随机放置障碍物）；'
        '不再依赖expert_density特征（首次移除专家轨迹密度依赖）；'
        '保存训练好的模型权重（.npz格式）。'
    )

    doc.add_heading('4.2 配置变化', level=2)
    add_styled_table(doc,
        ['参数', 'V3/V4', 'V5'],
        [
            ['训练样本', '43', '180'],
            ['候选路径/样本', '18', '18'],
            ['训练轮数', '220', '260'],
            ['学习率', '0.06', '0.035'],
            ['Train/Test分离', '按样本切分', '按demo隔离'],
            ['障碍物增强', '无', '训练+测试均开启'],
            ['expert_density', '保留', '移除'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('4.3 特征列表（11维）', level=2)
    doc.add_paragraph(
        '新增obstacle_clearance（障碍物间隙）和collision_free（无碰撞）特征，'
        '以及map_clearance替代原clearance特征。保留基础地图语义特征。'
    )

    doc.add_heading('4.4 分析', level=2)
    doc.add_paragraph(
        '虽然V5首次实现了train/test分离和障碍物感知，但180个样本太少，训练不充分。'
        '权重值域很小（0-0.2），说明训练收敛不足。多个特征权重收敛到0附近'
        '（drivable_ratio, collision_free, goal_reached），证明要学到有效的奖励函数需要更大的训练规模。'
    )

    # ── Chapter 5: V6 Heavy ──
    doc.add_heading('5. V6 Heavy：大规模自主障碍感知线性IRL', level=1)

    doc.add_heading('5.1 核心改进', level=2)
    doc.add_paragraph(
        '大幅度提升训练规模和数据复杂度。训练cases从180 → 2,400（+1233%），'
        '候选路径从18 → 24/case，每case固定7个障碍物，训练从260 → 900 epochs，'
        'batch size从全量 → 48（mini-batch SGD）。完全移除expert_density特征。'
    )

    doc.add_heading('5.2 配置对比', level=2)
    add_styled_table(doc,
        ['参数', 'V5', 'V6 Heavy'],
        [
            ['训练cases', '180', '2,400'],
            ['预处理耗时', '-', '30分钟'],
            ['候选路径/case', '18', '24'],
            ['障碍物/case', '若干', '7'],
            ['训练轮数', '260', '900'],
            ['学习率', '0.035', '0.035'],
            ['L2正则', '0.001', '0.0001'],
            ['Batch size', '全量', '48'],
            ['有效候选集', '-', '2,371'],
            ['平均候选数', '-', '16.65'],
            ['特征维度', '11', '12'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5.3 特征列表（12维）', level=2)
    add_styled_table(doc,
        ['特征', '含义', '特征', '含义'],
        [
            ['valid_ratio', '有效区域占比', 'obstacle_clearance_mean', '障碍物间隙均值'],
            ['lane_ratio', '车道区域占比', 'obstacle_clearance_min', '障碍物间隙最小值'],
            ['lane_connector_ratio', '连接区占比', 'goal_reached', '到达目标点'],
            ['non_intersection_ratio', '非交叉口占比', 'length_efficiency', '路径长度效率'],
            ['static_clearance_mean', '静态间隙均值', 'heading_to_goal', '朝向目标'],
            ['static_clearance_min', '静态间隙最小值', 'smoothness', '路径平滑度'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5.4 Loss收敛过程', level=2)
    add_styled_table(doc,
        ['轮数', 'NLL'],
        [
            ['0', '2.664'], ['100', '2.209'], ['200', '2.057'],
            ['300', '1.979'], ['400', '1.932'], ['500', '1.901'],
            ['600', '1.879'], ['700', '1.862'], ['800', '1.848'],
            ['899', '1.838'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Loss从2.664降到1.838，下降31.0%。')

    doc.add_heading('5.5 最终奖励权重', level=2)
    add_styled_table(doc,
        ['特征', '权重', '方向', '解读'],
        [
            ['length_efficiency', '+19.738', '正向', '偏好高效路径（最强信号）'],
            ['static_clearance_min', '+7.715', '正向', '偏好最小间隙大（安全）'],
            ['heading_to_goal', '+6.488', '正向', '偏好朝向目标'],
            ['obstacle_clearance_min', '-0.442', '负向', '统计偏差'],
            ['lane_ratio', '-4.069', '负向', '地图语义偏差'],
            ['obstacle_clearance_mean', '-5.481', '负向', '统计偏差'],
            ['smoothness', '-3.879', '负向', '候选生成偏见'],
            ['static_clearance_mean', '-15.369', '负向', '与min形成互补约束'],
            ['valid_ratio/lane_connector/goal_reached', '≈0', '-', '不作区分'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5.6 测试结果（30个复杂测试cases）', level=2)
    add_styled_table(doc,
        ['指标', 'IRL', 'Shortest Baseline', '对比'],
        [
            ['成功率', '100%', '100%', '持平'],
            ['平均路径长度', '133.85', '134.08', 'IRL略短'],
            ['最终距离误差', '0.0', '0.0', '持平'],
            ['最小静态间隙', '1.436', '1.436', '持平'],
            ['最小障碍物间隙', '4.67', '4.68', '持平'],
            ['障碍物违规', '0.0', '0.0', '持平'],
            ['地图违规', '0.0', '0.0', '持平'],
            ['平均转弯角', '0.0593', '0.0599', 'IRL略优'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('5.7 分析', level=2)
    bullets = [
        '规模效果显著：2,400 cases + 900 epochs使loss持续下降，未出现过拟合',
        '学到非平凡奖励结构：length_efficiency(+19.7)、static_clearance_min(+7.7)、heading_to_goal(+6.5)是主要正向驱动',
        'min比mean更有信息量：static_clearance_min正而mean负，说明模型关心最危险的点而非平均状态',
        '完全摆脱expert_density：模型不依赖专家轨迹密度特征，可泛化到新场景',
        '在简单障碍场景中，IRL与shortest baseline表现接近。复杂权衡场景（如窄通道+密集障碍物）需要非线性奖励函数',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # ── Chapter 6: V7 Neural ──
    doc.add_heading('6. V7 Neural Reward IRL（设计阶段）', level=1)

    doc.add_heading('6.1 设计目标', level=2)
    doc.add_paragraph(
        '从线性奖励升级为神经网络奖励函数，突破手工特征线性组合的表达能力上限。'
        '目标不是模仿专家路径，而是在没有专家路径输入时，给定start/goal/map/obstacles'
        '自主完成安全路径规划。'
    )

    doc.add_heading('6.2 模型架构', level=2)
    doc.add_paragraph(
        'MLP Reward Model: Input(18 features) → Linear(18,256) + ReLU + Dropout(0.10) '
        '→ Linear(256,256) + ReLU + Dropout(0.10) → Linear(256,1) → Scalar Reward Score'
    )
    doc.add_paragraph(
        '训练目标：Cross-Entropy Choice Loss —— 在14条候选路径中选择与teacher utility最匹配的路径。'
    )

    doc.add_heading('6.3 训练配置', level=2)
    add_styled_table(doc,
        ['参数', '值'],
        [
            ['数据集规模', '100,000 cases'],
            ['候选路径/case', '14'],
            ['障碍物/case', '8–24'],
            ['起点-目标距离', '55–210 cells'],
            ['Train/Val/Test', '70% / 15% / 15%'],
            ['训练轮数', '120'],
            ['Batch size', '768'],
            ['学习率', '2e-4'],
            ['优化器', 'AdamW + CosineAnnealingLR'],
            ['输入特征', '18维'],
        ]
    )
    doc.add_paragraph()

    # ── Chapter 7: Summary ──
    doc.add_heading('7. 版本对比总结', level=1)

    doc.add_heading('7.1 关键指标对比', level=2)
    add_styled_table(doc,
        ['版本', '核心改进', '训练规模', '特征维', '初始Loss', '最终Loss', 'IRL ADE'],
        [
            ['V3', 'Goal-conditioned基线', '43', '9', '2.052', '1.688', '2.600'],
            ['V4', '+Static Clearance', '43', '9', '2.052', '1.662', '1.919'],
            ['V5', '+Obstacle感知, Train/Test分离', '180', '11', '-', '-', '-'],
            ['V6 Heavy', '规模扩大, 完备特征', '2,400', '12', '2.664', '1.838', '-'],
            ['V7 Neural', '神经网络奖励', '100,000', '18→MLP', '-', '-', '-'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('7.2 方法论演进路径', level=2)
    doc.add_paragraph(
        'V3-V4: 手工特征探索（地图语义 + 专家密度 + 目标导向）\n'
        '→ V5: 训练框架化（Train/Test分离 + 障碍物增强 + 模型保存）\n'
        '→ V6 Heavy: 规模化训练（2400 cases + 900 epochs + 完备特征集）\n'
        '→ V7: 神经网络奖励（MLP替代线性θ + 100K大规模数据）'
    )

    doc.add_heading('7.3 核心发现', level=2)
    findings = [
        'Clearance特征是最有效的单点改进：V3→V4，加入static clearance使ADE下降26.2%',
        '规模决定质量：V5(180 samples)权重几乎为零，V6 Heavy(2400 cases)学到有意义的权重结构',
        'length_efficiency始终是最重要特征：在V3、V4、V6 Heavy中均获得最高权重',
        'min clearance比mean clearance更有信息量：V6 Heavy中static_clearance_min(+7.7) vs mean(-15.4)',
        '线性奖励在简单场景与shortest baseline差异不大，复杂场景需要神经网络',
    ]
    for i, f in enumerate(findings):
        doc.add_paragraph(f'{i+1}. {f}')

    # ── Chapter 8: Conclusion ──
    doc.add_heading('8. 结论与展望', level=1)

    doc.add_heading('8.1 主要结论', level=2)
    conclusions = [
        '最大熵IRL能够从nuPlan专家轨迹中学习有效的驾驶偏好，生成的地图合规路径成功率达100%',
        'Clearance特征（静态间隙、障碍物间隙）是连接地图感知与安全规划的关键桥梁',
        '大规模训练（2400+ cases）是线性IRL发挥作用的必要条件，小样本下权重无法收敛',
        '线性奖励在简单环境中与shortest baseline差异有限，需要神经网络来建模非线性驾驶偏好',
    ]
    for i, c in enumerate(conclusions):
        doc.add_paragraph(f'{i+1}. {c}')

    doc.add_heading('8.2 后续工作', level=2)
    future = [
        '完成V7 Neural Reward的100K数据集训练和测试评估',
        '引入闭环仿真评价（nuPlan simulator）',
        '探索Transformer/Attention结构建模路径级依赖',
        '引入交通规则约束（红绿灯、停止线、人行横道）',
        '多智能体交互场景下的IRL路径规划',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    # ── Appendix ──
    doc.add_heading('附录：实验环境', level=1)
    add_styled_table(doc,
        ['项目', '配置'],
        [
            ['硬件', 'NVIDIA Jetson (ARM aarch64, 61GiB RAM)'],
            ['OS', 'Ubuntu 20.04.6 LTS, L4T R35.4.1'],
            ['Python', '3.9.23 (Miniforge conda env: nuplan_irl)'],
            ['关键依赖', 'numpy, scipy, matplotlib, pandas, shapely, PyTorch'],
            ['数据', 'nuPlan v1.1 Mini Split + Maps v1.0'],
            ['地图', 'Las Vegas Strip (us-nv-las-vegas-strip)'],
        ]
    )
    doc.add_paragraph()

    # ── Insert figures ──
    doc.add_heading('附录：实验图表', level=1)

    loss_path = os.path.join(OUT_DIR, "report_loss_comparison.png")
    ade_path = os.path.join(OUT_DIR, "report_ade_comparison.png")
    weights_path = os.path.join(OUT_DIR, "report_v6_weights.png")

    if os.path.exists(loss_path):
        add_figure(doc, loss_path, '图1: 各版本训练Loss收敛对比')
        doc.add_paragraph()

    if os.path.exists(ade_path):
        add_figure(doc, ade_path, '图2: V3 vs V4 ADE对比')
        doc.add_paragraph()

    if os.path.exists(weights_path):
        add_figure(doc, weights_path, '图3: V6 Heavy学习的奖励权重')
        doc.add_paragraph()

    # Also include archives output figures if they exist
    v3_loss_path = "outputs/archive/nuplan_irl_goal_planning_v3/goal_planning_loss.png"
    v4_loss_path = "outputs/archive/nuplan_irl_goal_path_planning_v4/goal_planning_loss.png"
    v6_loss_path = "outputs/archive/nuplan_irl_autonomous_model_v6_heavy/autonomous_planning_loss.png"

    if os.path.exists(v6_loss_path):
        add_figure(doc, v6_loss_path, '图4: V6 Heavy完整训练Loss曲线（900 epochs）')

    # Save
    out_path = os.path.join(OUT_DIR, "IRL_实验报告_V3-V6.docx")
    doc.save(out_path)
    print(f"Report saved to: {out_path}")

# ── Main ───────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating figures...")
    make_figures()
    print("Building Word document...")
    build_docx()
    print("Done.")
